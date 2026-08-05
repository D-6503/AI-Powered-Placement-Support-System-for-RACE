import re
import os
import hashlib
import fitz  # PyMuPDF
import pdfplumber
import docx
from typing import Dict, List, Tuple, Any
from sqlalchemy.orm import Session
from backend.app.database import SkillsMaster, StudentSkill

def extract_text_from_pdf(file_path: str) -> str:
    text = ""
    # Try PyMuPDF first (fast and robust)
    try:
        doc = fitz.open(file_path)
        for page in doc:
            text += page.get_text()
        doc.close()
    except Exception as e:
        print(f"PyMuPDF error: {e}. Trying pdfplumber fallback.")
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e2:
            print(f"pdfplumber error: {e2}")
    return text

def extract_text_from_docx(file_path: str) -> str:
    text = ""
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        print(f"DOCX parser error: {e}")
    return text

def parse_contact_info(text: str) -> Dict[str, str]:
    contact = {"email": "", "phone": "", "linkedin": ""}
    
    # Email regex
    email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
    if email_match:
        contact["email"] = email_match.group(0)

    # Phone regex (supports international formats)
    phone_match = re.search(r'(\+?[\d\s\(\)-]{8,20})', text)
    if phone_match:
        # Filter potential noise (numbers should contain digits)
        phone_candidate = phone_match.group(0).strip()
        digits_count = sum(c.isdigit() for c in phone_candidate)
        if 7 <= digits_count <= 15:
            contact["phone"] = phone_candidate

    # LinkedIn regex
    linkedin_match = re.search(r'(linkedin\.com/in/[\w\-]+)', text, re.IGNORECASE)
    if linkedin_match:
        contact["linkedin"] = "https://" + linkedin_match.group(0)
        
    return contact

def split_resume_sections(text: str) -> Dict[str, str]:
    sections = {
        "contact_info": "",
        "skills": "",
        "experience": "",
        "projects": "",
        "education": ""
    }
    
    lines = text.split("\n")
    current_section = "contact_info"
    section_content = []
    
    # Section keywords map
    section_keywords = {
        "skills": ["skills", "technical skills", "technologies", "expertise", "core competencies", "technical expertise"],
        "experience": ["experience", "work experience", "employment history", "professional experience", "work history", "employment"],
        "projects": ["projects", "academic projects", "key projects", "personal projects", "engineering projects", "technical projects"],
        "education": ["education", "academic background", "credentials", "qualifications", "academic credentials"]
    }
    
    for line in lines:
        cleaned_line = line.strip().lower()
        if not cleaned_line:
            continue
            
        # Check if line indicates a new section
        found_new_section = False
        for sec_name, keywords in section_keywords.items():
            for kw in keywords:
                # Direct match or colon matches (e.g. "Skills:")
                if cleaned_line == kw or cleaned_line == f"{kw}:":
                    sections[current_section] += "\n".join(section_content)
                    current_section = sec_name
                    section_content = []
                    found_new_section = True
                    break
            if found_new_section:
                break
                
        if not found_new_section:
            section_content.append(line)
            
    # Save the last section
    sections[current_section] += "\n".join(section_content)
    
    # If some sections are empty, do a simpler regex fallback search
    for sec_name, keywords in section_keywords.items():
        if not sections[sec_name].strip():
            # Try finding the keyword and taking text until next major section heading
            for kw in keywords:
                pattern = re.compile(rf'(?:^|\n)({kw})(?:\n|:)', re.IGNORECASE)
                match = pattern.search(text)
                if match:
                    start_idx = match.end()
                    # Find end of this section (start of next keyword)
                    end_idx = len(text)
                    for other_sec, other_kws in section_keywords.items():
                        if other_sec == sec_name:
                            continue
                        for okw in other_kws:
                            other_pattern = re.compile(rf'(?:^|\n)({okw})(?:\n|:)', re.IGNORECASE)
                            other_match = other_pattern.search(text, start_idx)
                            if other_match and other_match.start() < end_idx:
                                end_idx = other_match.start()
                    sections[sec_name] = text[start_idx:end_idx].strip()
                    break
                    
    return sections

def extract_skills_from_text(text: str, db: Session) -> List[SkillsMaster]:
    matched_skills = []
    skills_catalog = db.query(SkillsMaster).all()
    
    text_lower = text.lower()
    
    for skill in skills_catalog:
        # Match by name
        name_esc = re.escape(skill.skill_name.lower())
        # Use word boundaries for short names, allow substring for longer technical terms
        if len(name_esc) <= 3:
            pattern = rf'\b{name_esc}\b'
        else:
            pattern = rf'\b{name_esc}\b'  # Still use word boundaries to avoid matching "java" in "javascript"
            
        if re.search(pattern, text_lower):
            matched_skills.append(skill)
            continue
            
        # Match by aliases
        if skill.alias_list:
            try:
                aliases = json.loads(skill.alias_list)
                for alias in aliases:
                    alias_esc = re.escape(alias.lower())
                    if len(alias_esc) <= 3:
                        p_alias = rf'\b{alias_esc}\b'
                    else:
                        p_alias = rf'\b{alias_esc}\b'
                        
                    if re.search(p_alias, text_lower):
                        matched_skills.append(skill)
                        break
            except Exception:
                pass
                
    return matched_skills

def evaluate_resume_quality(sections: Dict[str, str]) -> Tuple[float, List[Dict[str, Any]], List[str]]:
    checks = []
    improvements = []
    
    # 0. Basic Validation: Is it actually a resume?
    contact_text = (sections.get("contact_info", "") + " " + sections.get("education", "")).strip()
    email_found = bool(re.search(r'[\w\.-]+@[\w\.-]+\.\w+', contact_text))
    phone_found = bool(re.search(r'\+?[\d\s\(\)-]{8,20}', contact_text))
    
    has_skills = len(sections.get("skills", "").strip()) > 10
    has_experience = len(sections.get("experience", "").strip()) > 20
    
    if not (email_found or phone_found) and not (has_skills or has_experience):
        checks.append({
            "section": "Resume Verification",
            "status": False,
            "message": "Invalid Document: This PDF does not appear to contain standard resume fields or contact details."
        })
        improvements.append("Please upload a professional CV/Resume containing email, phone number, and sections for Skills and Work Experience.")
        return 0.0, checks, improvements
        
    score = 0.0
    
    # 1. Check sections presence (worth 60 points)
    section_weights = {
        "contact_info": 10.0,
        "skills": 15.0,
        "experience": 15.0,
        "projects": 10.0,
        "education": 10.0
    }
    
    for sec, weight in section_weights.items():
        content_len = len(sections[sec].strip())
        status = content_len > 30 if sec != "contact_info" else content_len > 10
        if status:
            score += weight
            checks.append({
                "section": sec.replace("_", " ").title(),
                "status": True,
                "message": f"Section '{sec}' is populated ({content_len} characters)."
            })
        else:
            checks.append({
                "section": sec.replace("_", " ").title(),
                "status": False,
                "message": f"Section '{sec}' is missing or too short."
            })
            improvements.append(f"Add detail or create a distinct section for '{sec.replace('_', ' ').title()}'.")
            
    # 2. Detail & length check (worth 25 points)
    total_len = sum(len(s.strip()) for s in sections.values())
    if total_len > 1500:
        score += 25
        checks.append({
            "section": "Overall Length",
            "status": True,
            "message": f"Resume is rich in content ({total_len} characters)."
        })
    elif total_len > 600:
        score += 15
        checks.append({
            "section": "Overall Length",
            "status": True,
            "message": "Resume length is moderate, but could use more detail."
        })
        improvements.append("Expand on your projects and work experience descriptions to add depth.")
    else:
        checks.append({
            "section": "Overall Length",
            "status": False,
            "message": "Resume is very short. Expand content."
        })
        improvements.append("Significantly expand resume content; add detailed projects, certifications, and skills.")
        
    # 3. Contact details specific checks (worth 15 points)
    contact_text = sections["contact_info"] + " " + sections["education"]
    email_found = bool(re.search(r'[\w\.-]+@[\w\.-]+\.\w+', contact_text))
    phone_found = bool(re.search(r'\+?[\d\s\(\)-]{8,20}', contact_text))
    linkedin_found = "linkedin.com" in contact_text.lower()
    
    contact_score = 0
    if email_found: contact_score += 5
    else: improvements.append("Add an email address to your contact information.")
    
    if phone_found: contact_score += 5
    else: improvements.append("Add a telephone number to your contact information.")
    
    if linkedin_found: contact_score += 5
    else: improvements.append("Include your LinkedIn profile link to improve online credibility.")
    
    score += contact_score
    checks.append({
        "section": "Contact Information Details",
        "status": contact_score == 15,
        "message": f"Found {sum([email_found, phone_found, linkedin_found])}/3 critical contact details."
    })
    
    return round(score, 1), checks, improvements
import json
